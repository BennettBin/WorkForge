from __future__ import annotations

from collections import Counter
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches
from openpyxl import load_workbook


@dataclass
class ExcelCateDistributionResult:
    cate_column: str
    total_rows: int
    categories: list[str]
    counts: list[int]
    chart_path: str | None
    report_path: str


class ExcelCateDistributionTool:
    def generate_report(
        self,
        *,
        excel_path: str,
        report_docx_path: str,
        chart_png_path: str,
        requirement: str,
        llm_markdown: str,
        language: str = "zh-CN",
    ) -> ExcelCateDistributionResult:
        excel_file = Path(excel_path)
        if not excel_file.exists():
            raise ValueError(f"Excel file not found: {excel_file}")

        wb = load_workbook(str(excel_file), data_only=True, read_only=True)
        try:
            ws = wb.active
            rows = list(ws.iter_rows(values_only=True))
        finally:
            wb.close()
        if not rows:
            raise ValueError("Excel sheet is empty.")

        header = [str(c).strip() if c is not None else "" for c in rows[0]]
        cate_idx = self._find_cate_column(header)
        if cate_idx is None:
            raise ValueError("Column `cate` not found in Excel header.")
        cate_col_name = header[cate_idx]

        values: list[str] = []
        for row in rows[1:]:
            if cate_idx >= len(row):
                continue
            cell = row[cate_idx]
            if cell is None:
                continue
            val = str(cell).strip()
            if val:
                values.append(val)
        if not values:
            raise ValueError("Column `cate` has no non-empty values.")

        counter = Counter(values)
        categories, counts = zip(*sorted(counter.items(), key=lambda x: (-x[1], x[0])))
        chart_path = self._save_bar_chart(list(categories), list(counts), chart_png_path, language=language)
        report_path = self._write_docx_report(
            report_docx_path=report_docx_path,
            chart_path=chart_path,
            cate_column=cate_col_name,
            total_rows=len(values),
            categories=list(categories),
            counts=list(counts),
            requirement=requirement,
            llm_markdown=llm_markdown,
            language=language,
        )
        return ExcelCateDistributionResult(
            cate_column=cate_col_name,
            total_rows=len(values),
            categories=list(categories),
            counts=list(counts),
            chart_path=chart_path,
            report_path=report_path,
        )

    def _find_cate_column(self, header: list[str]) -> int | None:
        lowered = [h.lower() for h in header]
        for i, name in enumerate(lowered):
            if name == "cate":
                return i
        for i, name in enumerate(lowered):
            if "cate" in name:
                return i
        return None

    def _save_bar_chart(self, categories: list[str], counts: list[int], chart_png_path: str, language: str) -> str | None:
        try:
            import matplotlib
            matplotlib.use("Agg")
            import matplotlib.pyplot as plt
        except Exception:
            return None

        fig, ax = plt.subplots(figsize=(10, 6), dpi=150)
        ax.bar(categories, counts, color="#4E79A7", edgecolor="#2F4B7C", linewidth=0.8)
        ax.set_xlabel("cate", fontsize=11)
        ax.set_ylabel("count", fontsize=11)
        if language.lower().startswith("zh"):
            ax.set_title("cate分布条形图", fontsize=13)
        else:
            ax.set_title("Category Distribution of `cate`", fontsize=13)
        ax.grid(axis="y", linestyle="--", alpha=0.35)
        ax.set_axisbelow(True)
        plt.xticks(rotation=30, ha="right")
        plt.tight_layout()
        out = Path(chart_png_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        fig.savefig(str(out))
        plt.close(fig)
        return str(out.resolve())

    def _write_docx_report(
        self,
        *,
        report_docx_path: str,
        chart_path: str | None,
        cate_column: str,
        total_rows: int,
        categories: list[str],
        counts: list[int],
        requirement: str,
        llm_markdown: str,
        language: str,
    ) -> str:
        doc = Document()
        title = "数据分析报告：cate分布统计" if language.lower().startswith("zh") else "Data Analysis Report: Distribution of `cate`"
        doc.add_heading(title, level=1)
        doc.add_paragraph(f"Requirement: {requirement}")
        doc.add_paragraph(f"Detected column: {cate_column}")
        doc.add_paragraph(f"Total valid rows: {total_rows}")

        doc.add_heading("Distribution Table", level=2)
        table = doc.add_table(rows=1, cols=3)
        header_cells = table.rows[0].cells
        header_cells[0].text = "rank"
        header_cells[1].text = "cate"
        header_cells[2].text = "count"
        for i, (cat, cnt) in enumerate(zip(categories, counts), start=1):
            cells = table.add_row().cells
            cells[0].text = str(i)
            cells[1].text = str(cat)
            cells[2].text = str(cnt)

        doc.add_heading("Bar Chart", level=2)
        if chart_path and Path(chart_path).exists():
            doc.add_picture(chart_path, width=Inches(6.6))
        else:
            doc.add_paragraph("Chart generation fallback: matplotlib unavailable in runtime environment.")

        doc.add_heading("Analysis", level=2)
        if llm_markdown.strip():
            doc.add_paragraph(llm_markdown[:6000])
        else:
            doc.add_paragraph("No model analysis text was generated.")

        report = Path(report_docx_path)
        report.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(report))
        return str(report.resolve())
